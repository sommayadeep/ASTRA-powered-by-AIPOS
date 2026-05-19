from __future__ import annotations

import io
import logging
from typing import Any
from typing import Optional

from fastapi import HTTPException, UploadFile

from backend.app.services.rag import index_document_chunks, list_indexed_documents

SUPPORTED_TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".csv"}
SUPPORTED_BINARY_EXTENSIONS = {".pdf", ".docx"}


def _get_extension(filename: str) -> str:
    dot_index = filename.rfind(".")
    return filename[dot_index:].lower() if dot_index != -1 else ""


def extract_text_from_upload(file: UploadFile) -> str:
    extension = _get_extension(file.filename or "")
    raw_bytes = file.file.read()

    if extension in SUPPORTED_TEXT_EXTENSIONS:
        return raw_bytes.decode("utf-8", errors="ignore")

    if extension == ".pdf":
        try:
            from pypdf import PdfReader
        except Exception as exc:  # pragma: no cover - dependency/runtime guard
            raise HTTPException(status_code=500, detail="PDF parsing is not available") from exc

        reader = PdfReader(io.BytesIO(raw_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)

    if extension == ".docx":
        try:
            from docx import Document
        except Exception as exc:  # pragma: no cover - dependency/runtime guard
            raise HTTPException(status_code=500, detail="DOCX parsing is not available") from exc

        document = Document(io.BytesIO(raw_bytes))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        return "\n".join(paragraphs)

    raise HTTPException(
        status_code=400,
        detail="Unsupported file type. Use PDF, DOCX, TXT, MD, or CSV.",
    )


def ingest_uploaded_document(file: UploadFile) -> dict[str, Any]:
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is required")

    if not file.content_type:
        raise HTTPException(status_code=400, detail="Content type is required")

    content = extract_text_from_upload(file)
    if not content.strip():
        raise HTTPException(status_code=400, detail="Uploaded file does not contain readable text")

    return index_document_chunks(
        filename=file.filename,
        content=content,
        content_type=file.content_type,
    )


def index_document_bytes(filename: str, raw_bytes: bytes, content_type: Optional[str] = None) -> Optional[dict[str, Any]]:
    logger = logging.getLogger(__name__)
    try:
        extension = _get_extension(filename or "")

        if extension in SUPPORTED_TEXT_EXTENSIONS:
            content = raw_bytes.decode("utf-8", errors="ignore")
        elif extension == ".pdf":
            try:
                from pypdf import PdfReader
            except Exception as exc:  # pragma: no cover - dependency/runtime guard
                logger.error("PDF parsing not available: %s", exc)
                return None
            reader = PdfReader(io.BytesIO(raw_bytes))
            pages = [page.extract_text() or "" for page in reader.pages]
            content = "\n".join(pages)
        elif extension == ".docx":
            try:
                from docx import Document
            except Exception as exc:  # pragma: no cover - dependency/runtime guard
                logger.error("DOCX parsing not available: %s", exc)
                return None
            document = Document(io.BytesIO(raw_bytes))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            content = "\n".join(paragraphs)
        else:
            logger.error("Unsupported file extension for background indexing: %s", extension)
            return None

        if not content.strip():
            logger.error("Uploaded file contains no readable text: %s", filename)
            return None

        return index_document_chunks(filename=filename, content=content, content_type=content_type)

    except Exception as exc:
        logger.exception("Error indexing uploaded document %s: %s", filename, exc)
        return None


def get_documents_overview() -> list[dict[str, Any]]:
    return list_indexed_documents()
